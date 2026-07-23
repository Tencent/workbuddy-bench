#!/usr/bin/env python3
from __future__ import annotations

import argparse
import csv
import json
import mimetypes
import re
import zipfile
from pathlib import Path


TEXT_SUFFIXES = {
    ".txt",
    ".md",
    ".markdown",
    ".json",
    ".jsonl",
    ".ndjson",
    ".csv",
    ".tsv",
    ".html",
    ".htm",
    ".svg",
    ".xml",
    ".yaml",
    ".yml",
    ".toml",
    ".ini",
    ".cfg",
    ".mmd",
    ".mermaid",
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".css",
    ".log",
    ".sh",
}

MARKDOWN_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*\S)\s*$", re.MULTILINE)
MARKDOWN_LINK_RE = re.compile(r"(?<!!)\[[^\]]*\]\(([^)]+)\)")
MARKDOWN_IMAGE_RE = re.compile(r"!\[[^\]]*\]\(([^)]+)\)")
MARKDOWN_CODE_FENCE_RE = re.compile(r"^```", re.MULTILINE)
MARKDOWN_TABLE_DIVIDER_RE = re.compile(r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?\s*$")


def describe_dir(path: Path, max_depth: int = 3) -> dict:
    entries = []
    root_depth = len(path.parts)
    for child in sorted(path.rglob("*")):
        rel_depth = len(child.parts) - root_depth
        if rel_depth > max_depth:
            continue
        rel = str(child.relative_to(path))
        entries.append(
            {
                "path": rel,
                "kind": "dir" if child.is_dir() else "file",
                "bytes": child.stat().st_size if child.is_file() else None,
            }
        )
    return {
        "path": str(path),
        "kind": "dir",
        "entries": entries,
    }


def load_docx(path: Path) -> dict:
    from docx import Document

    document = Document(path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)
    return {
        "path": str(path),
        "kind": "docx",
        "paragraph_count": len(paragraphs),
        "paragraphs": paragraphs[:200],
        "table_count": len(tables),
        "tables": tables[:20],
    }


def load_xlsx(path: Path) -> dict:
    from openpyxl import load_workbook

    workbook = load_workbook(path, data_only=True, read_only=True)
    sheets = []
    for worksheet in workbook.worksheets:
        rows = []
        for idx, row in enumerate(worksheet.iter_rows(values_only=True), start=1):
            rows.append(list(row))
            if idx >= 40:
                break
        sheets.append(
            {
                "title": worksheet.title,
                "preview_rows": rows,
            }
        )
    return {
        "path": str(path),
        "kind": "xlsx",
        "sheet_count": len(sheets),
        "sheets": sheets,
    }


def load_csv_preview(path: Path, limit: int = 40) -> dict:
    with path.open("r", encoding="utf-8", errors="replace", newline="") as fh:
        reader = csv.reader(fh)
        rows = []
        for index, row in enumerate(reader, start=1):
            rows.append(row)
            if index >= limit:
                break
    return {
        "path": str(path),
        "kind": "csv",
        "rows": rows,
    }


def load_zip_listing(path: Path) -> dict:
    with zipfile.ZipFile(path) as archive:
        names = sorted(archive.namelist())
    return {
        "path": str(path),
        "kind": "zip",
        "entries": names[:400],
    }


def load_text(path: Path, max_chars: int) -> str:
    text = path.read_text(encoding="utf-8", errors="replace")
    if len(text) > max_chars:
        half = max_chars // 2
        text = text[:half] + "\n...[truncated]...\n" + text[-half:]
    return text


def summarize_value(value, max_depth: int, max_items: int, max_string: int):
    if max_depth <= 0:
        if isinstance(value, dict):
            return {"_type": "dict", "_keys": list(value.keys())[:max_items], "_truncated": len(value) > max_items}
        if isinstance(value, list):
            return {"_type": "list", "_length": len(value), "_preview": value[:max_items]}
        if isinstance(value, str):
            return value[:max_string] + ("...[truncated]" if len(value) > max_string else "")
        return value

    if isinstance(value, dict):
        keys = list(value.keys())
        result = {}
        for key in keys[:max_items]:
            result[key] = summarize_value(value[key], max_depth - 1, max_items, max_string)
        if len(keys) > max_items:
            result["_truncated_keys"] = len(keys) - max_items
        return result

    if isinstance(value, list):
        result = [
            summarize_value(item, max_depth - 1, max_items, max_string)
            for item in value[:max_items]
        ]
        if len(value) > max_items:
            result.append({"_truncated_items": len(value) - max_items})
        return result

    if isinstance(value, str):
        return value[:max_string] + ("...[truncated]" if len(value) > max_string else "")

    return value


def is_local_reference(target: str) -> bool:
    value = (target or "").strip()
    if not value or value.startswith("#"):
        return False
    return not re.match(r"^[a-zA-Z][a-zA-Z0-9+.-]*:", value)


def count_markdown_table_blocks(lines: list[str]) -> int:
    count = 0
    for index in range(len(lines) - 1):
        if "|" not in lines[index]:
            continue
        if MARKDOWN_TABLE_DIVIDER_RE.match(lines[index + 1].strip()):
            count += 1
    return count


def summarize_markdown(path: Path, max_items: int, max_string: int) -> dict:
    text = path.read_text(encoding="utf-8", errors="replace")
    lines = text.splitlines()
    headings = []
    for match in MARKDOWN_HEADING_RE.finditer(text):
        headings.append(
            {
                "level": len(match.group(1)),
                "text": match.group(2)[:max_string],
            }
        )
        if len(headings) >= max_items:
            break

    image_targets = [
        target[:max_string]
        for target in MARKDOWN_IMAGE_RE.findall(text)
    ]
    link_targets = [
        target[:max_string]
        for target in MARKDOWN_LINK_RE.findall(text)
    ]
    local_targets = [
        target[:max_string]
        for target in (image_targets + link_targets)
        if is_local_reference(target)
    ]
    return {
        "path": str(path),
        "kind": "markdown_summary",
        "chars": len(text),
        "line_count": len(lines),
        "heading_count": len(MARKDOWN_HEADING_RE.findall(text)),
        "headings_preview": headings,
        "image_count": len(image_targets),
        "image_targets_preview": image_targets[:max_items],
        "link_count": len(link_targets),
        "link_targets_preview": link_targets[:max_items],
        "local_resource_count": len(local_targets),
        "local_resource_preview": local_targets[:max_items],
        "code_fence_count": len(MARKDOWN_CODE_FENCE_RE.findall(text)),
        "table_block_count": count_markdown_table_blocks(lines),
        "head_lines": [line[:max_string] for line in lines[:max_items]],
        "tail_lines": [line[:max_string] for line in lines[-max_items:]] if lines else [],
    }


def summarize_rule_evidence_summary(
    path: Path,
    payload: dict,
    max_items: int,
    max_string: int,
) -> dict:
    comparison = payload.get("comparison_highlights") or {}
    candidate = payload.get("candidate_item_assessment") or {}
    section_coverage = comparison.get("section_phrase_coverage") or {}
    required_images = comparison.get("required_image_paths") or {}
    code_anchor_presence = comparison.get("code_anchor_presence") or {}
    lead_phrase_match = comparison.get("lead_phrase_match") or {}
    forbidden_hits = comparison.get("forbidden_phrase_hits") or {}
    resource_paths = comparison.get("resource_paths_found") or {}
    raw_html = comparison.get("raw_html_outside_code") or {}

    simplified_candidate = {}
    if isinstance(candidate, dict):
        for item_id, item_value in list(candidate.items())[:max_items]:
            evidence = (item_value or {}).get("evidence") or {}
            simplified_candidate[item_id] = {
                "candidate_status": (item_value or {}).get("candidate_status"),
                "evidence": summarize_value(evidence, max_depth=1, max_items=max_items, max_string=max_string),
            }

    missing_sections = [
        section.get("expected_heading")
        for section in (section_coverage.get("sections") or [])
        if isinstance(section, dict) and not section.get("all_required_phrases_found")
    ][:max_items]

    return {
        "path": str(path),
        "kind": "rule_evidence_summary",
        "case_id": payload.get("case_id"),
        "eval_family": payload.get("eval_family"),
        "family_variant": payload.get("family_variant"),
        "artifact_exists": payload.get("artifact_exists"),
        "oracle_manifest_status": payload.get("oracle_manifest_status"),
        "completion_conditions": summarize_value(
            payload.get("completion_conditions") or [],
            max_depth=1,
            max_items=max_items,
            max_string=max_string,
        ),
        "candidate_item_assessment": simplified_candidate,
        "comparison_highlights": {
            "heading_preview": (comparison.get("heading_preview") or [])[:max_items],
            "ordered_headings_missing": (comparison.get("ordered_headings_missing") or [])[:max_items],
            "lead_phrase_match": {
                "matched_count": lead_phrase_match.get("matched_count"),
                "expected_count": lead_phrase_match.get("expected_count"),
                "missing": (lead_phrase_match.get("missing") or [])[:max_items],
            },
            "section_phrase_coverage": {
                "sections_full_match_count": section_coverage.get("sections_full_match_count"),
                "sections_total": section_coverage.get("sections_total"),
                "missing_sections": missing_sections,
            },
            "required_image_paths": {
                "expected_count": required_images.get("expected_count"),
                "missing": (required_images.get("missing") or [])[:max_items],
            },
            "code_anchor_presence": {
                "expected_count": code_anchor_presence.get("expected_count"),
                "missing": (code_anchor_presence.get("missing") or [])[:max_items],
            },
            "forbidden_phrase_hits": {
                "hit_count": forbidden_hits.get("hit_count"),
                "hits": (forbidden_hits.get("hits") or [])[:max_items],
            },
            "resource_paths_found": {
                "count": resource_paths.get("count"),
                "paths": (resource_paths.get("paths") or [])[:max_items],
            },
            "raw_html_outside_code": {
                "count": raw_html.get("count"),
                "sample_lines": (raw_html.get("sample_lines") or [])[:max_items],
            },
        },
    }


def summarize_html(path: Path, max_items: int, max_string: int) -> dict:
    from bs4 import BeautifulSoup

    text = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(text, "lxml")
    headings = []
    for node in soup.find_all(["h1", "h2", "h3"]):
        label = " ".join(node.get_text(" ", strip=True).split())
        if label:
            headings.append({"tag": node.name, "text": label[:max_string]})
        if len(headings) >= max_items:
            break
    paragraphs = []
    for node in soup.find_all("p"):
        label = " ".join(node.get_text(" ", strip=True).split())
        if label:
            paragraphs.append(label[:max_string])
        if len(paragraphs) >= max_items:
            break
    return {
        "path": str(path),
        "kind": "html_summary",
        "chars": len(text),
        "title": soup.title.get_text(strip=True)[:max_string] if soup.title else None,
        "meta_description": (
            (soup.find("meta", attrs={"name": "description"}) or {}).get("content", "")[:max_string]
            if soup.find("meta", attrs={"name": "description"})
            else None
        ),
        "heading_count": len(soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"])),
        "headings_preview": headings,
        "paragraph_preview": paragraphs,
        "image_count": len(soup.find_all("img")),
        "link_count": len(soup.find_all("a")),
        "table_count": len(soup.find_all("table")),
    }


def summarize_text(path: Path, max_items: int, max_string: int) -> dict:
    text = path.read_text(encoding="utf-8", errors="replace")
    lines = text.splitlines()
    return {
        "path": str(path),
        "kind": "text_summary",
        "chars": len(text),
        "line_count": len(lines),
        "head_lines": [line[:max_string] for line in lines[:max_items]],
        "tail_lines": [line[:max_string] for line in lines[-max_items:]] if lines else [],
    }


def json_error_summary(path: Path, exc: BaseException, max_string: int) -> dict:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except OSError:
        text = ""
    return {
        "path": str(path),
        "kind": "json_error",
        "error": f"{type(exc).__name__}: {exc}",
        "chars": len(text),
        "preview": text[:max_string],
    }


def summarize_file(path: Path, max_depth: int, max_items: int, max_string: int) -> dict:
    suffix = path.suffix.lower()
    if suffix == ".json":
        try:
            payload = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError) as exc:
            return json_error_summary(path, exc, max_string)
        if path.name == "rule-evidence-summary.json" and isinstance(payload, dict):
            return summarize_rule_evidence_summary(
                path,
                payload,
                max_items=max_items,
                max_string=max_string,
            )
        return {
            "path": str(path),
            "kind": "json_summary",
            "summary": summarize_value(
                payload,
                max_depth=max_depth,
                max_items=max_items,
                max_string=max_string,
            ),
        }
    if suffix in {".html", ".htm"}:
        return summarize_html(path, max_items=max_items, max_string=max_string)
    if suffix in {".md", ".markdown"}:
        return summarize_markdown(path, max_items=max_items, max_string=max_string)
    if suffix in TEXT_SUFFIXES:
        return summarize_text(path, max_items=max_items, max_string=max_string)
    payload = describe_file(path, max_chars=max_string)
    if isinstance(payload, str):
        return {
            "path": str(path),
            "kind": "text_summary",
            "chars": len(payload),
            "preview": payload[:max_string],
        }
    return payload


def summarize_dir(path: Path, max_items: int) -> dict:
    entries = []
    for child in sorted(path.rglob("*")):
        if len(entries) >= max_items:
            break
        rel = str(child.relative_to(path))
        entries.append(
            {
                "path": rel,
                "kind": "dir" if child.is_dir() else "file",
                "bytes": child.stat().st_size if child.is_file() else None,
            }
        )
    return {
        "path": str(path),
        "kind": "dir_summary",
        "entry_count_preview": len(entries),
        "entries": entries,
    }


def describe_file(path: Path, max_chars: int) -> dict | str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return load_docx(path)
    if suffix == ".xlsx":
        return load_xlsx(path)
    if suffix in {".csv", ".tsv"}:
        return load_csv_preview(path)
    if suffix == ".zip":
        return load_zip_listing(path)
    if suffix == ".json":
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError) as exc:
            return json_error_summary(path, exc, max_chars)
    if suffix in TEXT_SUFFIXES:
        return load_text(path, max_chars)

    mime_type, _ = mimetypes.guess_type(path.name)
    return {
        "path": str(path),
        "kind": "binary",
        "mime_type": mime_type or "application/octet-stream",
        "bytes": path.stat().st_size,
    }


def cmd_show(path: Path, max_chars: int):
    if not path.exists():
        raise SystemExit(f"path not found: {path}")
    payload = describe_dir(path) if path.is_dir() else describe_file(path, max_chars)
    if isinstance(payload, str):
        print(payload)
    else:
        print(json.dumps(payload, indent=2, ensure_ascii=False))


def cmd_summary(path: Path, max_depth: int, max_items: int, max_string: int):
    if not path.exists():
        raise SystemExit(f"path not found: {path}")
    payload = (
        summarize_dir(path, max_items=max_items)
        if path.is_dir()
        else summarize_file(path, max_depth=max_depth, max_items=max_items, max_string=max_string)
    )
    print(json.dumps(payload, indent=2, ensure_ascii=False))


def main():
    parser = argparse.ArgumentParser(description="Inspect runtime verification artifacts.")
    subparsers = parser.add_subparsers(dest="command", required=True)

    show = subparsers.add_parser("show", help="Show a structured preview of a file or directory.")
    show.add_argument("path", type=Path)
    show.add_argument("--max-chars", type=int, default=12000)

    summary = subparsers.add_parser("summary", help="Show a compact summary instead of full content.")
    summary.add_argument("path", type=Path)
    summary.add_argument("--max-depth", type=int, default=2)
    summary.add_argument("--max-items", type=int, default=8)
    summary.add_argument("--max-string", type=int, default=240)

    args = parser.parse_args()
    if args.command == "show":
        cmd_show(args.path.resolve(), args.max_chars)
    elif args.command == "summary":
        cmd_summary(
            args.path.resolve(),
            max_depth=args.max_depth,
            max_items=args.max_items,
            max_string=args.max_string,
        )


if __name__ == "__main__":
    main()
